import os
from django.http import HttpResponse
from docx import Document
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def convert_days(timetables,time):
    timetables_days = []
    for i in timetables:
        timetable=[]
        for idx,j in enumerate(i):
            days={}
            days.update({'time':time[idx],'M': j[0],"T": j[1],'W': j[2],"Th": j[3],"F": j[4]})
            timetable.append(days)
        timetables_days.append(timetable)

    return timetables_days


def am_pm(time):
    if int(time)<12:
        time=time+".00am"
    else:
        time=time+".00pm"
    return time

def generate_docx(timetable,data):
    document = Document()

    sections = document.sections
        
    sec_pr = document.sections[0]._sectPr # get the section properties el
    # create new borders el
    pg_borders = OxmlElement('w:pgBorders')
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section



    paragraph=document.add_heading("M.H. Saboo Siddik College of Engineering",level=1)
    font = paragraph.style.font
    font.size=Pt(17)
    font.underline = True
    paragraph.alignment=1


    paragraph=document.add_paragraph(f'Deaprtment of {data.get("department")}', style='Intense Quote')
    paragraph.alignment=1


    #header table
    htable=document.add_table(1,3)

    htab_cells=htable.rows[0].cells

    ht0=htab_cells[0].add_paragraph()
    ht0.add_run(f'Room:{data.get("rooms")}')
    # fontstyle(ht0,font_bold=True)
    ht0.alignment = 0

    ht1=htab_cells[1].add_paragraph()
    ht1.add_run(f"Semester: {data.get('semester')}")

    ht1.alignment = 1

    ht2=htab_cells[2].add_paragraph()
    ht2.add_run(f"Date :{data.get('date')}")

    ht2.alignment = 2

    # main time table
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Day/Time'
    hdr_cells[1].text = 'Monday'
    hdr_cells[2].text = 'Tuesday'
    hdr_cells[3].text = 'Wednesday'
    hdr_cells[4].text = 'Thursday'
    hdr_cells[5].text = 'Friday'

    for records in timetable:
        row_cells = table.add_row().cells
        row_cells[0].text = records['time']
        row_cells[1].text = records['M']
        row_cells[2].text = records['T']
        row_cells[3].text = records['W']
        row_cells[4].text = records['Th']
        row_cells[5].text = records['F']

    document.add_page_break()

    path = "media"
    # Check whether the specified path exists or not
    isExist = os.path.exists(path)
    if not isExist:
    # Create a new directory because it does not exist
        os.makedirs(path)

    document.save(f"media/{data.get('semester')}.docx")
    

